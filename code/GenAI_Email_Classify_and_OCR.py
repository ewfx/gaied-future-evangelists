#### Team Name Future_Evangelists ####
#### Team Members: Venkata Evani, Srinivasu Choppa and Anil Devarala ####
#### \nWelcome to the Email Classification and OCR tool.
#### This tool can read and interpret email file (with or without attachment) provided as input or document. Supports .eml,.pdf, .docs and .txt formats
#### Then it will classify the input and identify Request Type, Sub Request Type and Reasoning
#### This program will run multiple times until we provide 'exit' as input.
    



import os
import email
from email import policy
from email.parser import BytesParser
import docx
import PyPDF2
import extract_msg
import google.generativeai as genai
import sys
from pathlib import Path
from docx import Document


#setting GenAI Google Cloud APIKEY and Model details

genai.configure(api_key="AIzaSyDOdTi-H6E6u720iPExpIOtRUpFW2wdrVs")

model = genai.GenerativeModel("gemini-1.5-pro-002")

def extract_text_from_attachment(filename, payload):
    ext = os.path.splitext(filename)[1].lower()
    content = ""

    try:
        if ext == ".txt":
            content = payload.get_payload(decode=True).decode('utf-8', errors='ignore')

        elif ext == ".pdf":
            from io import BytesIO
            reader = PyPDF2.PdfReader(BytesIO(payload.get_payload(decode=True)))
            for page in reader.pages:
                content += page.extract_text() or ""

        elif ext == ".docx":
            from io import BytesIO
            doc = docx.Document(BytesIO(payload.get_payload(decode=True)))
            content = "\n".join([p.text for p in doc.paragraphs])

        else:
            content = f"[Unsupported attachment type: {ext}]"

    except Exception as e:
        content = f"[Error reading {filename}: {e}]"

    return content

def read_eml_with_attachments(file_path):
    with open(file_path, 'rb') as f:
        msg = BytesParser(policy=policy.default).parse(f)

    email_body = ""
    attachments = {}

    if msg.is_multipart():
        for part in msg.walk():
            content_disposition = part.get("Content-Disposition", "")
            content_type = part.get_content_type()

            if "attachment" in content_disposition:
                filename = part.get_filename()
                if filename:
                    attachments[filename] = part

            elif content_type == "text/plain" and not email_body:
                email_body = part.get_payload(decode=True).decode('utf-8', errors='ignore')
    else:
        email_body = msg.get_payload(decode=True).decode('utf-8', errors='ignore')

    return email_body, attachments

def read_pdf(file_path):
    content = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                content += page.extract_text() or ""
    except Exception as e:
        content = f"[Error reading PDF: {e}]"
    return content

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"[Error reading DOCX: {e}]"

def read_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"[Error reading TXT file: {e}]"

def detect_file_type(file_path):
    return os.path.splitext(file_path)[1].lower()

def interactive_input_reader(file_path):

    content = ""
    if not os.path.exists(file_path):
        print("❌ File not found.")
        return

    ext = detect_file_type(file_path)

    file_type_map = {
        ".eml": "Email File (.eml)",
        ".pdf": "PDF Document (.pdf)",
        ".docx": "Word Document (.docx)",
        ".txt": "Text File (.txt)"
    }

    file_type_description = file_type_map.get(ext, "Unknown File Type")
    print(f"\n📂 Detected file type: {file_type_description}")

    if ext == ".eml":
        print("\nReading email...\n")
        email_body, attachments = read_eml_with_attachments(file_path)
        #print("📧 Email Body:\n")
        #print(email_body)
        #print("\n---\n")
        content = email_body
        if attachments:
            print(f"📎 Found {len(attachments)} attachment(s):")
            for i, fname in enumerate(attachments.keys(), 1):
                print(f"  {i}. {fname}")

            choice = input("\nEnter the number of the attachment to view (or 'all'): ").strip()

            if choice.lower() == 'all':
                for fname, part in attachments.items():
                    print(f"\n📄 {fname}:\n")
                    print(extract_text_from_attachment(fname, part))
            else:
                try:
                    idx = int(choice) - 1
                    fname = list(attachments.keys())[idx]
                    print(f"\n📄 {fname}:\n")
                    print(extract_text_from_attachment(fname, attachments[fname]))
                    priority = input("Do you want to include this attachment content in the email body? (y/n): ").strip().lower()
                    if priority == 'y':
                        content += extract_text_from_attachment(fname, attachments[fname]) # Append attachment content to the email body
                    print(content)
                except:
                    print("⚠️ Invalid selection.")
        else:
            print("No attachments found.")

    elif ext == ".pdf":
        content = read_pdf(file_path)
        #print("\n📄 PDF Content:\n")
        #print(content)

    elif ext == ".docx":
        content = read_docx(file_path)
        #print("\n📄 DOCX Content:\n")
        #print(content)

    elif ext == ".txt":
        content = read_txt(file_path)
        #print("\n📄 TXT Content:\n")
        #print(content)

    else:
        print("❌ Unsupported file type.")

    try:
    
        header_content = """
Classify the following email into one of the request types:
[Adjustment, AU Transfer, Closing Notice, Commitment Change, Fee Payment, Money Movement-Inbound, Money Movement - Outbound, Urgent, Acknowledgement].
and Request subtypes:
[Reallocation Fees,Amendement Fees,Reallocation Principal,Cashless Roll,Decrease,Increase,Ongoing Fee,Letter of Credit Fee, Principal,Interest, 
Principal + interest, Principal+interest+fee,Timebound, foreign currency]"""
        tailer_content = """"         
Return the following:
- Request Type:
- Request Subtype:
- Sender:
- Receiver(To):
- Date:
- Amount:
- Currency:
- Reasoning:"""

        prompt = header_content + content + tailer_content
        print("Classifying the extracted content:")
        #print(prompt)

    
        response = model.generate_content(prompt)
        print(response.text)
    except Exception as e:
        print(f"Error: {e}")

os.system('cls' if os.name == 'nt' else 'clear')
print("\nWelcome to the Email Classification and OCR tool.")
print("This tool can read and interpret email file provided as input or document. Supports .eml,.pdf, .docs and .txt formats")
print("Then it will classify the input and identify Request Type, Sub Request Type and Reasoning")
print("This program will run multiple times until we provide 'exit' as input.")
print("  ")
print(" ")

while True:
    
    print("  ")
    UNDERLINE = "\033[4m"
    RESET = "\033[0m"

    
    file_path = input(f"{UNDERLINE}Enter the path to your file (.eml, .pdf, .docx, .txt), (type exit to quit):{RESET}  ").strip()

    
    if file_path.lower() == 'exit':
        print("Exiting the program.")
        break
    

    # Start the interactive session
    interactive_input_reader(file_path)